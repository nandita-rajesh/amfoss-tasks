import json
from typing import Final
import csv
from io import StringIO, BytesIO
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
import requests
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes, CallbackQueryHandler

TOKEN: Final = '7523879986:AAGRJgc9zulWSKhMAiumYTiqROkUK0xGcCg'
BOT_USERNAME: Final = '@PagePalBot_bot'
GOOGLE_BOOKS_API_KEY: Final = 'AIzaSyDu22IbjtybOsBuhebZLyebqBDKt8CxJPs'

# Fetch books from Google Books API
def fetch_books(query: str):
    URL = f'https://www.googleapis.com/books/v1/volumes?q={query}&key={GOOGLE_BOOKS_API_KEY}'
    response = requests.get(URL)
    print(f"API Response Status Code: {response.status_code}")
    print(f"API Response: {response.text}") 
    if response.status_code == 200:
        data = response.json()
        books = []
        for item in data.get('items', []):
            volume_info = item.get('volumeInfo', {})
            book = {
                "title": volume_info.get('title', 'No title'),
                "author": ', '.join(volume_info.get('authors', [])),
                "description": volume_info.get('description', 'No description'),
                "year": volume_info.get('publishedDate', 'No date').split('-')[0],
                "language": volume_info.get('language', 'No language'),
                "preview": volume_info.get('previewLink', 'No link')
            }
            books.append(book)
        return books
    else:
        return []

# Add a book to the reading list
def add_to_reading_list(book_title, preview_link):
    book_title_lower = book_title.lower()
    try:
        doc = Document("reading_list.docx")
    except Exception as e:
        if isinstance(e, PackageNotFoundError):
            doc = Document()
            doc.add_heading('Reading List', level=1)
        else:
            return f"Error reading the document: {str(e)}"

    book_found = False
    for paragraph in doc.paragraphs:
        if paragraph.text.lower().startswith(f"book title: {book_title_lower}"):
            book_found = True
            break

    if book_found:
        return "Book is already in the reading list."

    doc.add_paragraph(f"Book Title: {book_title}")
    doc.add_paragraph(f"Preview Link: {preview_link}")
    doc.save("reading_list.docx")
    return None

# Delete a book from the reading list
def delete_from_reading_list(book_title):
    book_title_lower = book_title.lower()
    try:
        doc = Document("reading_list.docx")
    except Exception as e:
        if isinstance(e, PackageNotFoundError):
            return "Reading list is empty. No book to delete."
        else:
            return f"Error reading the document: {str(e)}"

    new_doc = Document()
    new_doc.add_heading('Reading List', level=1)
    book_found = False

    paragraphs_to_keep = []
    skip_next = False

    for paragraph in doc.paragraphs:
        text = paragraph.text.lower()
        if text.startswith(f"book title: {book_title_lower}"):
            book_found = True
            skip_next = True
            continue  
        if skip_next and text.startswith("preview link:"):
            skip_next = False 
            continue  
        paragraphs_to_keep.append(paragraph.text)

    for paragraph_text in paragraphs_to_keep:
        new_doc.add_paragraph(paragraph_text)

    if book_found:
        new_doc.save("reading_list.docx")
        return None 
    else:
        return "Book not found."

# Command handlers
async def start_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Hello! I am PagePal, your book recommendation assistant!")

async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    help_text = (
        "/start - Welcome message.\n"
        "/book - Ask for a book genre and return a list of books.\n"
        "/preview - Ask for a book name and return the preview link.\n"
        "/list - Manage your reading list.\n"
        "/reading_list - Manage your reading list with buttons.\n"
        "/help - Show this help message."
    )
    await update.message.reply_text(help_text)

async def book_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Please type the genre of the book you are interested in.")
    context.user_data['awaiting_genre'] = True

async def preview_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Please type the name of the book you want a preview link for.")
    context.user_data['awaiting_book_name'] = True

async def list_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Please use /reading_list to manage your reading list.")

async def reading_list_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    keyboard = [
        [InlineKeyboardButton("Add a book", callback_data='add_book')],
        [InlineKeyboardButton("Delete a book", callback_data='delete_book')],
        [InlineKeyboardButton("View Reading List", callback_data='view_list')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text("Manage your reading list:", reply_markup=reply_markup)

async def add_book_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Please type the name of the book you want to add.")
    context.user_data['awaiting_add'] = True

async def delete_book_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Please type the book title you want to delete.")
    context.user_data['awaiting_delete'] = True

# Messages handlers
async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip()

    if context.user_data.get('awaiting_genre'):
        genre = text.lower()
        books = fetch_books(f'subject:{genre}')
        if books:
            output = StringIO()
            writer = csv.writer(output)
            writer.writerow(["Title", "Author", "Description", "Year", "Language", "Preview"])
            for book in books:
                writer.writerow([book["title"], book["author"], book["description"], book["year"], book["language"], book["preview"]])
            output.seek(0)
            await update.message.reply_document(document=output, filename=f"{genre}_books.csv")
        else:
            await update.message.reply_text("No books found for this genre. Please try again.")
        context.user_data.pop('awaiting_genre', None)
        return

    if context.user_data.get('awaiting_book_name'):
        book_name = text.lower()
        books = fetch_books(f'title:{text}')
        for book in books:
            if book["title"].lower() == book_name:
                await update.message.reply_text(f"Preview link for '{text}': {book['preview']}")
                context.user_data.pop('awaiting_book_name', None)
                return
        await update.message.reply_text("Book not found. Please try again.")
        context.user_data.pop('awaiting_book_name', None)
        return

    if context.user_data.get('awaiting_add'):
        books = fetch_books(f'title:{text}')
        for book in books:
            if book["title"].lower() == text.lower():
                result = add_to_reading_list(book["title"], book["preview"])
                if result:
                    await update.message.reply_text(result)
                else:
                    await update.message.reply_text(f"Book '{text}' added to reading list.")
                context.user_data.pop('awaiting_add', None)
                return
        await update.message.reply_text("Book not found. Please try adding another book.")
        context.user_data.pop('awaiting_add', None)
        return

    if context.user_data.get('awaiting_delete'):
        result = delete_from_reading_list(text)
        if result:
            await update.message.reply_text(result)
        else:
            await update.message.reply_text(f"Book '{text}' deleted from reading list.")
        context.user_data.pop('awaiting_delete', None)
        return

    await update.message.reply_text("Command not recognized. Use /help for a list of commands.")

# Handle button interactions in the reading list
async def button(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    data = query.data

    if data == 'add_book':
        await query.message.reply_text("To add a book, use /add_book command.")
    elif data == 'delete_book':
        await query.message.reply_text("To delete a book, use /delete_book command.")
    elif data == 'view_list':
        try:
            doc = Document("reading_list.docx")
        except Exception:
            doc = Document()
            doc.add_heading('Reading List', level=1)
        
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        await query.message.reply_document(document=output, filename="reading_list.docx")

# Error handler
async def error(update: Update, context: ContextTypes.DEFAULT_TYPE):
    print(f"Update {update} caused error {context.error}")

# Main function to start the bot
if __name__ == '__main__':
    print('Starting bot...')
    app = Application.builder().token(TOKEN).build()

    app.add_handler(CommandHandler('start', start_command))
    app.add_handler(CommandHandler('help', help_command))
    app.add_handler(CommandHandler('book', book_command))
    app.add_handler(CommandHandler('preview', preview_command))
    app.add_handler(CommandHandler('list', list_command))
    app.add_handler(CommandHandler('reading_list', reading_list_command))
    app.add_handler(CommandHandler('add_book', add_book_command))
    app.add_handler(CommandHandler('delete_book', delete_book_command))
    app.add_handler(MessageHandler(filters.TEXT, handle_message))
    app.add_handler(CallbackQueryHandler(button))

    print('Bot started!')
    app.run_polling(poll_interval=3)